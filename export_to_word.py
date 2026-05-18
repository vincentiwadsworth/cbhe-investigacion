"""Convert govtake-sudamerica/paper.md to Word (.docx)"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Read markdown
with open("govtake-sudamerica/paper.md", encoding="utf-8") as f:
    md = f.read()

lines = md.split("\n")

style_map = {
    "# ": "Heading 1",
    "## ": "Heading 2",
    "### ": "Heading 3",
}

in_table = False
table_rows = []

for line in lines:
    # Skip YAML frontmatter metadata blockquote
    if line.startswith("> "):
        p = doc.add_paragraph(line[2:])
        p.style = doc.styles["Normal"]
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        continue

    # Detect headings
    is_heading = False
    for prefix, style_name in style_map.items():
        if line.startswith(prefix):
            doc.add_heading(line[len(prefix):], level=int(prefix.count("#")))
            is_heading = True
            break
    if is_heading:
        continue

    # Horizontal rule
    if line.strip() == "---":
        doc.add_paragraph("_" * 60)
        continue

    # Table detection
    if "|" in line and line.strip().startswith("|"):
        if not in_table:
            in_table = True
            table_rows = []
        # Skip separator rows (|---|---|)
        if not re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            table_rows.append(cells)
        continue
    else:
        if in_table and table_rows:
            # Flush table
            num_cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = "Light Grid Accent 1"
            for i, row in enumerate(table_rows):
                for j, cell_text in enumerate(row):
                    if j < num_cols:
                        cell = table.rows[i].cells[j]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
            doc.add_paragraph("")  # spacing
            in_table = False
            table_rows = []
            continue

    # Bold text (**text**)
    if line.strip():
        p = doc.add_paragraph()
        # Handle inline formatting
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", line)
        for part in parts:
            run = p.add_run(part.replace("*", ""))
            run.font.size = Pt(10)
            if part.startswith("**") and part.endswith("**"):
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run.italic = True
    else:
        doc.add_paragraph("")

doc.save("govtake-sudamerica/paper.docx")
print("✅ paper.docx saved")
